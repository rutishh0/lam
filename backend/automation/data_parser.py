"""
Universal Data Parser
Parses user data from various file formats (CSV, PDF, DOC, TXT, MD)
"""

import logging
import json
import csv
import io
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import re

# Optional imports for advanced file formats
try:
    import pypdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import markdown
    MD_SUPPORT = True
except ImportError:
    MD_SUPPORT = False

logger = logging.getLogger(__name__)

class DataParser:
    """Parse structured data from various file formats"""
    
    def __init__(self):
        self.field_patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b'),
            'zip': re.compile(r'\b[0-9]{5}(?:-[0-9]{4})?\b'),
            'date': re.compile(r'\b(?:\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\d{4}[-/]\d{1,2}[-/]\d{1,2})\b'),
            'ssn': re.compile(r'\b[0-9]{3}-[0-9]{2}-[0-9]{4}\b'),
            'url': re.compile(r'https?://(?:www\.)?[a-zA-Z0-9-]+(?:\.[a-zA-Z]{2,})+(?:/[^\s]*)?')
        }
    
    async def parse_file(self, file_content: Union[str, bytes], file_type: str, filename: str = '') -> List[Dict[str, Any]]:
        """
        Parse data from various file formats
        
        Args:
            file_content: File content as string or bytes
            file_type: File extension (csv, pdf, docx, txt, md)
            filename: Original filename
            
        Returns:
            List of dictionaries containing parsed data
        """
        file_type = file_type.lower().strip('.')
        
        try:
            if file_type == 'csv':
                return await self._parse_csv(file_content)
            elif file_type == 'txt':
                return await self._parse_txt(file_content)
            elif file_type == 'pdf' and PDF_SUPPORT:
                if isinstance(file_content, str):
                    file_content = file_content.encode('utf-8')
                return await self._parse_pdf(file_content)
            elif file_type in ['doc', 'docx'] and DOCX_SUPPORT:
                if isinstance(file_content, str):
                    file_content = file_content.encode('utf-8')
                return await self._parse_docx(file_content)
            elif file_type == 'md' and MD_SUPPORT:
                return await self._parse_markdown(file_content)
            elif file_type == 'json':
                return await self._parse_json(file_content)
            else:
                # Try to parse as text
                return await self._parse_txt(file_content)
                
        except Exception as e:
            logger.error(f"Error parsing {file_type} file: {str(e)}")
            raise Exception(f"Failed to parse {file_type} file: {str(e)}")
    
    async def _parse_csv(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Parse CSV file content"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        results = []
        csv_reader = csv.DictReader(io.StringIO(content))
        
        for row in csv_reader:
            # Clean and normalize the data
            cleaned_row = {}
            for key, value in row.items():
                if key and value:
                    # Normalize key names
                    normalized_key = key.lower().strip().replace(' ', '_').replace('-', '_')
                    cleaned_row[normalized_key] = value.strip()
            
            if cleaned_row:
                # Try to infer data types and enhance data
                enhanced_row = await self._enhance_data(cleaned_row)
                results.append(enhanced_row)
        
        return results
    
    async def _parse_txt(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Parse plain text file using intelligent extraction"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        results = []
        
        # Try to detect if it's structured data
        lines = content.strip().split('\n')
        
        # Check if it looks like key-value pairs
        if self._is_key_value_format(lines):
            data = await self._parse_key_value(lines)
            if data:
                results.append(data)
        else:
            # Extract data using patterns and heuristics
            extracted_data = await self._extract_from_text(content)
            if extracted_data:
                results.append(extracted_data)
        
        return results
    
    async def _parse_pdf(self, content: bytes) -> List[Dict[str, Any]]:
        """Parse PDF file content"""
        if not PDF_SUPPORT:
            raise Exception("PDF support not available. Install pypdf: pip install pypdf")
        
        try:
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            text_content = ""
            
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            # Parse the extracted text
            return await self._parse_txt(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
    
    async def _parse_docx(self, content: bytes) -> List[Dict[str, Any]]:
        """Parse DOCX file content"""
        if not DOCX_SUPPORT:
            raise Exception("DOCX support not available. Install python-docx: pip install python-docx")
        
        try:
            doc = docx.Document(io.BytesIO(content))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"
            
            # Parse the extracted text
            return await self._parse_txt(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    async def _parse_markdown(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Parse Markdown file content"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        if MD_SUPPORT:
            # Convert markdown to HTML then extract text
            html_content = markdown.markdown(content)
            # Simple HTML tag removal
            text_content = re.sub('<[^<]+?>', '', html_content)
        else:
            # Just use the raw markdown text
            text_content = content
        
        return await self._parse_txt(text_content)
    
    async def _parse_json(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Parse JSON file content"""
        if isinstance(content, bytes):
            content = content.decode('utf-8', errors='ignore')
        
        try:
            data = json.loads(content)
            
            # If it's already a list of dicts, return it
            if isinstance(data, list):
                return [await self._enhance_data(item) if isinstance(item, dict) else {'data': item} for item in data]
            elif isinstance(data, dict):
                return [await self._enhance_data(data)]
            else:
                return [{'data': data}]
                
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON: {str(e)}")
            raise Exception(f"Invalid JSON format: {str(e)}")
    
    def _is_key_value_format(self, lines: List[str]) -> bool:
        """Check if text appears to be in key-value format"""
        key_value_count = 0
        
        for line in lines[:10]:  # Check first 10 lines
            if ':' in line or '=' in line:
                parts = re.split('[:|=]', line, 1)
                if len(parts) == 2 and len(parts[0].strip()) < 50:
                    key_value_count += 1
        
        return key_value_count >= 3
    
    async def _parse_key_value(self, lines: List[str]) -> Dict[str, Any]:
        """Parse key-value formatted text"""
        data = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try different separators
            for separator in [':', '=', '-']:
                if separator in line:
                    parts = line.split(separator, 1)
                    if len(parts) == 2:
                        key = parts[0].strip().lower().replace(' ', '_').replace('-', '_')
                        value = parts[1].strip()
                        if key and value:
                            data[key] = value
                        break
        
        return await self._enhance_data(data)
    
    async def _extract_from_text(self, text: str) -> Dict[str, Any]:
        """Extract structured data from unstructured text"""
        data = {}
        
        # Extract emails
        emails = self.field_patterns['email'].findall(text)
        if emails:
            data['email'] = emails[0]
        
        # Extract phone numbers
        phones = self.field_patterns['phone'].findall(text)
        if phones:
            phone = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
            data['phone'] = phone
        
        # Extract names (heuristic approach)
        name_match = re.search(r'(?:Name|Contact|Person)[:\s]+([A-Za-z\s]+)', text, re.I)
        if name_match:
            data['full_name'] = name_match.group(1).strip()
        
        # Extract address-like content
        address_match = re.search(r'(?:Address|Location)[:\s]+(.+?)(?:\n|$)', text, re.I)
        if address_match:
            data['address'] = address_match.group(1).strip()
        
        # Extract zip codes
        zips = self.field_patterns['zip'].findall(text)
        if zips:
            data['zip_code'] = zips[0]
        
        # Extract dates
        dates = self.field_patterns['date'].findall(text)
        if dates:
            # Try to identify birth dates
            for i, date in enumerate(dates):
                if re.search(r'(?:birth|dob|born)', text[:text.find(date)].lower()):
                    data['date_of_birth'] = date
                    break
        
        # Extract company names
        company_match = re.search(r'(?:Company|Organization|Employer)[:\s]+([A-Za-z0-9\s&.,-]+)', text, re.I)
        if company_match:
            data['company'] = company_match.group(1).strip()
        
        # Extract URLs
        urls = self.field_patterns['url'].findall(text)
        if urls:
            data['website'] = urls[0]
        
        return await self._enhance_data(data)
    
    async def _enhance_data(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Enhance and validate extracted data"""
        enhanced = data.copy()
        
        # Split full name into first and last if needed
        if 'full_name' in enhanced and 'first_name' not in enhanced:
            name_parts = enhanced['full_name'].split()
            if len(name_parts) >= 2:
                enhanced['first_name'] = name_parts[0]
                enhanced['last_name'] = ' '.join(name_parts[1:])
        
        # Combine first and last name if full name is missing
        if 'first_name' in enhanced and 'last_name' in enhanced and 'full_name' not in enhanced:
            enhanced['full_name'] = f"{enhanced['first_name']} {enhanced['last_name']}"
        
        # Normalize phone numbers
        if 'phone' in enhanced:
            phone = re.sub(r'\D', '', enhanced['phone'])
            if len(phone) == 10:
                enhanced['phone'] = f"{phone[:3]}-{phone[3:6]}-{phone[6:]}"
            elif len(phone) == 11 and phone[0] == '1':
                enhanced['phone'] = f"{phone[1:4]}-{phone[4:7]}-{phone[7:]}"
        
        # Validate email
        if 'email' in enhanced:
            email = enhanced['email'].lower()
            if self.field_patterns['email'].match(email):
                enhanced['email'] = email
            else:
                del enhanced['email']
        
        # Add metadata
        enhanced['_parsed_at'] = datetime.utcnow().isoformat()
        enhanced['_confidence'] = self._calculate_confidence(enhanced)
        
        return enhanced
    
    def _calculate_confidence(self, data: Dict[str, Any]) -> float:
        """Calculate confidence score for parsed data"""
        confidence = 0.0
        important_fields = ['email', 'first_name', 'last_name', 'phone', 'address']
        
        for field in important_fields:
            if field in data and data[field]:
                confidence += 0.2
        
        # Additional confidence for validated fields
        if 'email' in data and self.field_patterns['email'].match(data['email']):
            confidence += 0.1
        
        if 'phone' in data and len(re.sub(r'\D', '', data['phone'])) >= 10:
            confidence += 0.1
        
        return min(confidence, 1.0) 